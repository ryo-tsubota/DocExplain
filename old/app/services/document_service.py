"""
Document processing service
"""
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional
import asyncio
import logging

from fastapi import HTTPException, UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
    UnstructuredMarkdownLoader
)
from sentence_transformers import SentenceTransformer
import pypdf
import docx
import markdown

from app.core.config import settings
from app.models.document import (
    Document, 
    DocumentChunk,
    DocumentResponse, 
    DocumentSummary, 
    ProcessingStatus, 
    SupportedFileType,
    generate_content_hash
)

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for handling document operations"""
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        # In-memory storage for demo - in production, use a database
        self._documents: dict[str, Document] = {}
        self._document_chunks: dict[str, List[DocumentChunk]] = {}
        
        # Initialize text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize embedding model
        self.embedding_model = None
        self._initialize_embedding_model()
    
    async def upload_document(self, file: UploadFile) -> DocumentResponse:
        """
        Upload and process a document file
        
        Args:
            file: The uploaded file
            
        Returns:
            DocumentResponse with document metadata
            
        Raises:
            HTTPException: If file validation fails or processing errors occur
        """
        # Validate file
        await self._validate_file(file)
        
        # Read file content
        content = await file.read()
        await file.seek(0)  # Reset file pointer
        
        # Generate document metadata
        document_id = str(uuid.uuid4())
        content_hash = generate_content_hash(content)
        file_type = self._get_file_type(file.filename)
        
        # Save file to disk
        file_path = self.upload_dir / f"{document_id}_{file.filename}"
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Create document record
        document = Document(
            id=document_id,
            filename=file.filename,
            file_type=file_type,
            upload_timestamp=datetime.now(timezone.utc),
            processing_status=ProcessingStatus.PENDING,
            content_hash=content_hash,
            chunk_count=0,
            file_size=len(content)
        )
        
        # Store document metadata
        self._documents[document_id] = document
        
        # Start processing in background
        asyncio.create_task(self._process_document_async(document_id, file_path))
        
        return DocumentResponse(
            document=document,
            message="Document uploaded successfully and queued for processing"
        )
    
    async def get_documents(self) -> List[DocumentSummary]:
        """Get list of all uploaded documents"""
        return [
            DocumentSummary(
                id=doc.id,
                filename=doc.filename,
                file_type=doc.file_type,
                upload_timestamp=doc.upload_timestamp,
                processing_status=doc.processing_status,
                chunk_count=doc.chunk_count,
                file_size=doc.file_size
            )
            for doc in self._documents.values()
        ]
    
    async def get_document(self, document_id: str) -> Optional[Document]:
        """Get a specific document by ID"""
        return self._documents.get(document_id)
    
    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and its associated files
        
        Args:
            document_id: The document ID to delete
            
        Returns:
            True if deleted successfully, False if document not found
        """
        document = self._documents.get(document_id)
        if not document:
            return False
        
        # Remove file from disk
        file_path = self.upload_dir / f"{document_id}_{document.filename}"
        if file_path.exists():
            file_path.unlink()
        
        # Remove chunks from memory
        if document_id in self._document_chunks:
            del self._document_chunks[document_id]
        
        # Remove from memory storage
        del self._documents[document_id]
        return True
    
    async def get_processing_status(self, document_id: str) -> Optional[ProcessingStatus]:
        """Get the processing status of a document"""
        document = self._documents.get(document_id)
        return document.processing_status if document else None
    
    async def _validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file
        
        Args:
            file: The uploaded file to validate
            
        Raises:
            HTTPException: If validation fails
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in settings.allowed_file_types:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_ext}. Allowed types: {', '.join(settings.allowed_file_types)}"
            )
        
        # Check file size
        if file.size and file.size > settings.max_file_size:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds maximum allowed size of {settings.max_file_size / (1024*1024):.1f}MB"
            )
        
        # Validate content type if available
        if file.content_type:
            allowed_content_types = {
                'application/pdf': '.pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                'text/plain': '.txt',
                'text/markdown': '.md'
            }
            
            expected_ext = allowed_content_types.get(file.content_type)
            if expected_ext and file_ext != expected_ext:
                raise HTTPException(
                    status_code=400,
                    detail=f"Content type {file.content_type} does not match file extension {file_ext}"
                )
    
    def _get_file_type(self, filename: str) -> SupportedFileType:
        """Determine file type from filename"""
        ext = Path(filename).suffix.lower()
        type_mapping = {
            '.pdf': SupportedFileType.PDF,
            '.docx': SupportedFileType.DOCX,
            '.txt': SupportedFileType.TXT,
            '.md': SupportedFileType.MD
        }
        return type_mapping[ext]
    
    def _initialize_embedding_model(self) -> None:
        """Initialize the sentence transformer model for embeddings"""
        try:
            # Use a lightweight but effective model for embeddings
            model_name = getattr(settings, 'embedding_model', 'all-MiniLM-L6-v2')
            self.embedding_model = SentenceTransformer(model_name)
            logger.info(f"Initialized embedding model: {model_name}")
        except Exception as e:
            logger.error(f"Failed to initialize embedding model: {e}")
            self.embedding_model = None
    
    async def _process_document_async(self, document_id: str, file_path: Path) -> None:
        """
        Process document asynchronously - extract text, chunk, and generate embeddings
        
        Args:
            document_id: The document ID to process
            file_path: Path to the uploaded file
        """
        try:
            # Update status to processing
            if document_id in self._documents:
                self._documents[document_id].processing_status = ProcessingStatus.PROCESSING
            
            # Extract text from document
            text_content = await self._extract_text(file_path)
            
            if not text_content.strip():
                raise ValueError("No text content could be extracted from the document")
            
            # Create chunks with overlap
            chunks = await self._create_chunks(document_id, text_content)
            
            # Generate embeddings for chunks
            if self.embedding_model:
                await self._generate_embeddings(chunks)
            
            # Store chunks
            self._document_chunks[document_id] = chunks
            
            # Update document status and chunk count
            if document_id in self._documents:
                self._documents[document_id].processing_status = ProcessingStatus.COMPLETED
                self._documents[document_id].chunk_count = len(chunks)
            
            logger.info(f"Successfully processed document {document_id} with {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Failed to process document {document_id}: {e}")
            if document_id in self._documents:
                self._documents[document_id].processing_status = ProcessingStatus.FAILED
    
    async def _extract_text(self, file_path: Path) -> str:
        """
        Extract text content from various document formats
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If text extraction fails
        """
        try:
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                return await self._extract_docx_text(file_path)
            elif file_ext == '.txt':
                return await self._extract_txt_text(file_path)
            elif file_ext == '.md':
                return await self._extract_markdown_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise ValueError(f"Failed to extract text from document: {e}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            loader = PyPDFLoader(str(file_path))
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        except Exception as e:
            # Fallback to pypdf directly
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    return "\n\n".join(text_content)
            except Exception as fallback_e:
                raise ValueError(f"PDF text extraction failed: {e}, fallback also failed: {fallback_e}")
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            loader = Docx2txtLoader(str(file_path))
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        except Exception as e:
            # Fallback to python-docx directly
            try:
                doc = docx.Document(file_path)
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                return "\n\n".join(paragraphs)
            except Exception as fallback_e:
                raise ValueError(f"DOCX text extraction failed: {e}, fallback also failed: {fallback_e}")
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            loader = TextLoader(str(file_path), encoding='utf-8')
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        except Exception as e:
            # Fallback to direct file reading
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except Exception as fallback_e:
                raise ValueError(f"TXT text extraction failed: {e}, fallback also failed: {fallback_e}")
    
    async def _extract_markdown_text(self, file_path: Path) -> str:
        """Extract text from Markdown file"""
        try:
            loader = UnstructuredMarkdownLoader(str(file_path))
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        except Exception as e:
            # Fallback to direct markdown processing
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    md_content = file.read()
                    # Convert markdown to plain text
                    html = markdown.markdown(md_content)
                    # Simple HTML tag removal (basic approach)
                    import re
                    text = re.sub('<[^<]+?>', '', html)
                    return text
            except Exception as fallback_e:
                raise ValueError(f"Markdown text extraction failed: {e}, fallback also failed: {fallback_e}")
    
    async def _create_chunks(self, document_id: str, text_content: str) -> List[DocumentChunk]:
        """
        Create overlapping chunks from text content
        
        Args:
            document_id: The document ID
            text_content: The extracted text content
            
        Returns:
            List of DocumentChunk objects
        """
        try:
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text_content)
            
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    id=f"{document_id}_chunk_{i}",
                    document_id=document_id,
                    content=chunk_text,
                    chunk_index=i,
                    metadata={
                        "chunk_size": len(chunk_text),
                        "total_chunks": len(text_chunks)
                    }
                )
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Chunking failed for document {document_id}: {e}")
            raise ValueError(f"Failed to create chunks: {e}")
    
    async def _generate_embeddings(self, chunks: List[DocumentChunk]) -> None:
        """
        Generate embeddings for document chunks
        
        Args:
            chunks: List of DocumentChunk objects to generate embeddings for
        """
        if not self.embedding_model:
            logger.warning("Embedding model not available, skipping embedding generation")
            return
        
        try:
            # Extract text content from chunks
            texts = [chunk.content for chunk in chunks]
            
            # Generate embeddings in batch
            embeddings = self.embedding_model.encode(texts, convert_to_tensor=False)
            
            # Assign embeddings to chunks
            for chunk, embedding in zip(chunks, embeddings):
                # Handle both numpy arrays and lists
                if hasattr(embedding, 'tolist'):
                    chunk.embedding = embedding.tolist()
                else:
                    chunk.embedding = list(embedding)
                chunk.metadata["embedding_model"] = self.embedding_model.get_sentence_embedding_dimension()
            
            logger.info(f"Generated embeddings for {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            # Don't raise exception - embeddings are optional for basic functionality
    
    async def get_document_chunks(self, document_id: str) -> List[DocumentChunk]:
        """
        Get all chunks for a specific document
        
        Args:
            document_id: The document ID
            
        Returns:
            List of DocumentChunk objects
        """
        return self._document_chunks.get(document_id, [])
    
    async def search_similar_chunks(self, query: str, document_ids: Optional[List[str]] = None, top_k: int = 5) -> List[DocumentChunk]:
        """
        Search for similar chunks using semantic similarity
        
        Args:
            query: The search query
            document_ids: Optional list of document IDs to search within
            top_k: Number of top results to return
            
        Returns:
            List of most similar DocumentChunk objects
        """
        if not self.embedding_model:
            logger.warning("Embedding model not available, cannot perform semantic search")
            return []
        
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query], convert_to_tensor=False)[0]
            
            # Collect all chunks to search
            all_chunks = []
            if document_ids:
                for doc_id in document_ids:
                    all_chunks.extend(self._document_chunks.get(doc_id, []))
            else:
                for chunks in self._document_chunks.values():
                    all_chunks.extend(chunks)
            
            # Filter chunks that have embeddings
            chunks_with_embeddings = [chunk for chunk in all_chunks if chunk.embedding]
            
            if not chunks_with_embeddings:
                return []
            
            # Calculate similarities
            from numpy import dot
            from numpy.linalg import norm
            
            similarities = []
            for chunk in chunks_with_embeddings:
                # Calculate cosine similarity
                similarity = dot(query_embedding, chunk.embedding) / (norm(query_embedding) * norm(chunk.embedding))
                similarities.append((chunk, similarity))
            
            # Sort by similarity and return top_k
            similarities.sort(key=lambda x: x[1], reverse=True)
            return [chunk for chunk, _ in similarities[:top_k]]
            
        except Exception as e:
            logger.error(f"Semantic search failed: {e}")
            return []


# Global service instance
document_service = DocumentService()